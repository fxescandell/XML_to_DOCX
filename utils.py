import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import json
import os

CONFIG_FILE = 'styles_config.json'

def load_config():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, 'r') as f:
                return json.load(f)
        except json.JSONDecodeError:
            return {}
    else:
        return {}

def save_config(config):
    with open(CONFIG_FILE, 'w') as f:
        json.dump(config, f, indent=4)
    print(f"Configuración guardada en {CONFIG_FILE}")

def apply_styles(paragraph, text, style_name, style_type, doc):
    styles = doc.styles
    if style_name not in [style.name for style in styles]:
        if style_type == 'parrafo':
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        elif style_type == 'caracter':
            style = styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
        font = style.font
        font.size = Pt(12)  # Ajusta el tamaño de fuente si es necesario

    if style_type == 'parrafo':
        paragraph.style = style_name
        paragraph.add_run(text)
    elif style_type == 'caracter':
        run = paragraph.add_run(text)
        run.style = style_name

def clean_default_styles(doc):
    styles = doc.styles
    keep_styles = {s['style'] for s in load_config().values()}
    keep_styles.add("Agenda-General-Parrafo")  # Include the general paragraph style

    for style in list(styles):
        if style.type in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER) and style.name not in keep_styles:
            styles.element.remove(style.element)

def process_combined_elements(elements, field1, field2, paragraph, config, doc):
    text1 = elements[0].text.strip() if elements[0] is not None and elements[0].text else ""
    text2 = elements[1].text.strip() if elements[1] is not None and elements[1].text else ""

    if text1:
        style_name1 = config.get(field1, {}).get('style', "Agenda-General-Parrafo")
        style_type1 = config.get(field1, {}).get('type', 'caracter')  # Default to 'caracter'
        apply_styles(paragraph, text1, style_name1, style_type1, doc)

    if text1 and text2:
        paragraph.add_run(" · ")

    if text2:
        style_name2 = config.get(field2, {}).get('style', "Agenda-General-Parrafo")
        style_type2 = config.get(field2, {}).get('type', 'caracter')  # Default to 'caracter'
        apply_styles(paragraph, text2, style_name2, style_type2, doc)

def process_fields(parent_element, fields, doc, config):
    processed_fields = set()
    for field in fields:
        if field in ["Evento-Principal-Hora", "Sub-evento-Hora", "actividad-hora"]:
            if field not in processed_fields:
                paragraph = None
                if field == "Evento-Principal-Hora":
                    if parent_element.find('Evento-Principal-Hora') is not None or parent_element.find('Evento-Principal-Lugar') is not None:
                        paragraph = doc.add_paragraph()
                        process_combined_elements([parent_element.find('Evento-Principal-Hora'), parent_element.find('Evento-Principal-Lugar')], 'Evento-Principal-Hora', 'Evento-Principal-Lugar', paragraph, config, doc)
                        processed_fields.update(['Evento-Principal-Hora', 'Evento-Principal-Lugar'])
                elif field == "Sub-evento-Hora":
                    if parent_element.find('Sub-evento-Hora') is not None or parent_element.find('Sub-evento-Lugar') is not None:
                        paragraph = doc.add_paragraph()
                        process_combined_elements([parent_element.find('Sub-evento-Hora'), parent_element.find('Sub-evento-Lugar')], 'Sub-evento-Hora', 'Sub-evento-Lugar', paragraph, config, doc)
                        processed_fields.update(['Sub-evento-Hora', 'Sub-evento-Lugar'])
                elif field == "actividad-hora":
                    if parent_element.find('actividad-hora') is not None or parent_element.find('actividad-lugar') is not None:
                        paragraph = doc.add_paragraph()
                        process_combined_elements([parent_element.find('actividad-hora'), parent_element.find('actividad-lugar')], 'actividad-hora', 'actividad-lugar', paragraph, config, doc)
                        processed_fields.update(['actividad-hora', 'actividad-lugar'])
        else:
            if field not in processed_fields:
                element = parent_element.find(field)
                if element is not None and element.text and element.text.strip():
                    paragraph = doc.add_paragraph()
                    style_name = config.get(field, {}).get('style', "Agenda-General-Parrafo")
                    style_type = config.get(field, {}).get('type', 'parrafo')  # Default to 'parrafo'
                    apply_styles(paragraph, element.text.strip(), style_name, style_type, doc)
                    processed_fields.add(field)

def process_xml_to_docx(xml_file, output_folder, output_file_name):
    config = load_config()
    tree = ET.parse(xml_file)
    root = tree.getroot()

    doc = Document()

    # Create "Agenda-General-Parrafo" style
    if "Agenda-General-Parrafo" not in [style.name for style in doc.styles]:
        general_style = doc.styles.add_style("Agenda-General-Parrafo", WD_STYLE_TYPE.PARAGRAPH)
        general_style.font.size = Pt(12)

    for event in root.findall('Evento-Principal'):
        process_fields(event, config.keys(), doc, config)

        programa = event.find('Evento-Principal-Programa')
        if programa is not None:
            for sub_event in programa.findall('Sub-evento'):
                process_fields(sub_event, config.keys(), doc, config)

                actividades = sub_event.find('Sub-evento-actividades')
                if actividades is not None:
                    for actividad in actividades.findall('actividad'):
                        process_fields(actividad, config.keys(), doc, config)

    # Clean default styles
    clean_default_styles(doc)

    # Apply "Agenda-General-Parrafo" style to all paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.style is None or paragraph.style.name == 'Normal':  # Change default Normal style to Agenda-General-Parrafo
            paragraph.style = "Agenda-General-Parrafo"

    output_path = os.path.join(output_folder, output_file_name)
    doc.save(output_path)
    print(f"Document saved to {output_path}")
